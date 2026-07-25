"""Turn an uploaded file into a list of (page_number, text) pairs.

Keeping page numbers all the way through ingestion is what lets the UI cite
"MDR, Article 10(4), p. 45" instead of just naming the file.
"""

from __future__ import annotations

from pathlib import Path

Page = tuple[int, str]


class UnsupportedFileType(Exception):
    pass


def read_pdf(path: Path) -> list[Page]:
    import fitz  # PyMuPDF

    pages: list[Page] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text and text.strip():
                pages.append((index, text))
    return pages


def read_docx(path: Path) -> list[Page]:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    # python-docx exposes no page boundaries, so the document is a single "page".
    return [(1, "\n".join(parts))] if parts else []


def read_text(path: Path) -> list[Page]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [(1, text)] if text.strip() else []


READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".txt": read_text,
    ".md": read_text,
}


def read_document(path: Path) -> list[Page]:
    reader = READERS.get(path.suffix.lower())
    if reader is None:
        raise UnsupportedFileType(f"Unsupported file type: {path.suffix}")
    return reader(path)
